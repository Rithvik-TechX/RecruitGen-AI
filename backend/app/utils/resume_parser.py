"""
Resume Text Extraction Utilities.

Supports PDF (PyMuPDF / fitz) and DOCX (python-docx) formats.
"""

from __future__ import annotations

import os

import structlog

logger = structlog.get_logger(__name__)


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text content from a PDF file using PyMuPDF."""
    import fitz  # PyMuPDF

    try:
        parts: list[str] = []
        with fitz.open(file_path) as doc:
            for page in doc:
                page_text = page.get_text("text")
                if page_text.strip():
                    parts.append(page_text)
                else:
                    # Fallback: extract from text blocks
                    for block in page.get_text("blocks"):
                        if block[6] == 0:  # text block type
                            parts.append(block[4])
        full_text = "\n".join(parts).strip()
        logger.info(
            "pdf_text_extracted",
            file=os.path.basename(file_path),
            chars=len(full_text),
        )
        return full_text
    except Exception as exc:
        logger.error("pdf_extraction_failed", file=file_path, error=str(exc))
        raise ValueError(f"Failed to parse PDF file: {exc}") from exc


def extract_text_from_docx(file_path: str) -> str:
    """Extract text content from a DOCX file using python-docx."""
    from docx import Document

    try:
        doc = Document(file_path)
        parts: list[str] = []

        # Paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    parts.append(" | ".join(row_text))

        full_text = "\n".join(parts).strip()
        logger.info(
            "docx_text_extracted",
            file=os.path.basename(file_path),
            chars=len(full_text),
        )
        return full_text
    except Exception as exc:
        logger.error("docx_extraction_failed", file=file_path, error=str(exc))
        raise ValueError(f"Failed to parse DOCX file: {exc}") from exc


def extract_text_from_resume(file_path: str) -> str:
    """Extract text from a resume file based on its extension."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported: .pdf, .docx")
